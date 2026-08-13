import io

from docx import Document
from docx.shared import Inches

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as ExcelImage


def build_word_report(
    reviewed_results,
    review_style,
    uploaded_screens
):
    """
    Build a downloadable Word UX review report.
    """

    document = Document()

    document.add_heading(
        "UX Microcopy Review",
        level=0
    )

    document.add_paragraph(
        f"UX Style Guide: {review_style}"
    )

    #
    # Calculate summary
    #

    screens_reviewed = len(
        reviewed_results
    )

    microcopy_reviewed = 0

    acceptable_count = 0
    improvement_count = 0
    issue_count = 0

    for result in reviewed_results:

        microcopy_items = result.get(
            "microcopy_items",
            []
        )

        microcopy_reviewed += len(
            microcopy_items
        )

        for finding in result.get(
            "findings",
            []
        ):

            assessment = (
                finding.get(
                    "assessment",
                    ""
                ).lower()
            )

            if assessment == "acceptable":

                acceptable_count += 1

            elif assessment == "improvement":

                improvement_count += 1

            elif assessment == "issue":

                issue_count += 1

    #
    # Summary
    #

    document.add_heading(
        "Review Summary",
        level=1
    )

    summary = document.add_table(
        rows=0,
        cols=2
    )

    summary.style = "Table Grid"

    summary_rows = [

        (
            "Style Guide",
            review_style
        ),

        (
            "Screens Reviewed",
            str(screens_reviewed)
        ),

        (
            "Microcopy Items Reviewed",
            str(microcopy_reviewed)
        ),

        (
            "Acceptable",
            str(acceptable_count)
        ),

        (
            "Improvements",
            str(improvement_count)
        ),

        (
            "Issues",
            str(issue_count)
        )

    ]

    for label, value in summary_rows:

        cells = summary.add_row().cells

        cells[0].text = label
        cells[1].text = value

    #
    # Screen reports
    #

    for result in reviewed_results:

        document.add_page_break()

        screen_name = result.get(
            "screen",
            "Screen"
        )

        document.add_heading(
            screen_name,
            level=1
        )

        #
        # Embed screenshot
        #

        screen_file = next(

            (
                screen
                for screen in uploaded_screens
                if screen.name == screen_name
            ),

            None

        )

        if screen_file:

            image_bytes = (
                screen_file.getvalue()
            )

            image_stream = io.BytesIO(
                image_bytes
            )

            try:

                document.add_picture(
                    image_stream,
                    width=Inches(6.5)
                )

            except Exception:

                document.add_paragraph(
                    "Screenshot could not be "
                    "embedded in the report."
                )

        #
        # Findings
        #

        findings = result.get(
            "findings",
            []
        )

        if not findings:

            document.add_paragraph(
                "No UX writing findings were "
                "identified on this screen."
            )

            continue

        document.add_heading(
            "UX Findings",
            level=2
        )

        for index, finding in enumerate(
            findings,
            start=1
        ):

            document.add_heading(

                f"{index}. "
                f"{finding.get('text', '')}",

                level=3

            )

            document.add_paragraph(

                f"Assessment: "
                f"{finding.get('assessment', '')}"

            )

            document.add_paragraph(

                f"Element: "
                f"{finding.get('element_type', '')}"

            )

            document.add_paragraph(

                f"Context: "
                f"{finding.get('approximate_context', '')}"

            )

            recommended_text = (
                finding.get(
                    "recommended_text"
                )
            )

            if recommended_text:

                document.add_paragraph(

                    f"Recommended microcopy: "
                    f"{recommended_text}"

                )

            document.add_paragraph(

                f"Reason: "
                f"{finding.get('reason', '')}"

            )

            document.add_paragraph(

                f"Style rule: "
                f"{finding.get('style_rule', '')}"

            )

            document.add_paragraph(

                f"Severity: "
                f"{finding.get('severity', '')}"

            )

    #
    # Return downloadable bytes
    #

    output = io.BytesIO()

    document.save(
        output
    )

    output.seek(0)

    return output.getvalue()


def build_excel_report(
    reviewed_results,
    review_style
):
    """
    Build a downloadable Excel UX review report.
    """

    workbook = Workbook()

    #
    # Summary sheet
    #

    summary_sheet = (
        workbook.active
    )

    summary_sheet.title = "Summary"

    screens_reviewed = len(
        reviewed_results
    )

    microcopy_reviewed = 0

    acceptable_count = 0
    improvement_count = 0
    issue_count = 0

    for result in reviewed_results:

        microcopy_reviewed += len(
            result.get(
                "microcopy_items",
                []
            )
        )

        for finding in result.get(
            "findings",
            []
        ):

            assessment = (
                finding.get(
                    "assessment",
                    ""
                ).lower()
            )

            if assessment == "acceptable":

                acceptable_count += 1

            elif assessment == "improvement":

                improvement_count += 1

            elif assessment == "issue":

                issue_count += 1

    summary_sheet["A1"] = (
        "UX Microcopy Review"
    )

    summary_sheet["A1"].font = Font(
        bold=True,
        size=16
    )

    summary_sheet["A3"] = "Style Guide"
    summary_sheet["B3"] = review_style

    summary_sheet["A4"] = "Screens Reviewed"
    summary_sheet["B4"] = screens_reviewed

    summary_sheet["A5"] = (
        "Microcopy Items Reviewed"
    )

    summary_sheet["B5"] = (
        microcopy_reviewed
    )

    summary_sheet["A6"] = "Acceptable"
    summary_sheet["B6"] = acceptable_count

    summary_sheet["A7"] = "Improvements"
    summary_sheet["B7"] = improvement_count

    summary_sheet["A8"] = "Issues"
    summary_sheet["B8"] = issue_count

    summary_sheet.column_dimensions[
        "A"
    ].width = 30

    summary_sheet.column_dimensions[
        "B"
    ].width = 45

    #
    # Findings sheet
    #

    findings_sheet = (
        workbook.create_sheet(
            "Findings"
        )
    )

    headers = [

        "Screen",

        "Element",

        "Context",

        "Current Copy",

        "Assessment",

        "Recommendation",

        "Reason",

        "Style Rule",

        "Severity"

    ]

    for column_index, header in enumerate(
        headers,
        start=1
    ):

        cell = (
            findings_sheet.cell(
                row=1,
                column=column_index
            )
        )

        cell.value = header

        cell.font = Font(
            bold=True
        )

        cell.alignment = (
            Alignment(
                wrap_text=True
            )
        )

    row = 2

    for result in reviewed_results:

        screen_name = result.get(
            "screen",
            ""
        )

        for finding in result.get(
            "findings",
            []
        ):

            values = [

                screen_name,

                finding.get(
                    "element_type",
                    ""
                ),

                finding.get(
                    "approximate_context",
                    ""
                ),

                finding.get(
                    "text",
                    ""
                ),

                finding.get(
                    "assessment",
                    ""
                ),

                finding.get(
                    "recommended_text",
                    ""
                ),

                finding.get(
                    "reason",
                    ""
                ),

                finding.get(
                    "style_rule",
                    ""
                ),

                finding.get(
                    "severity",
                    ""
                )

            ]

            for column_index, value in enumerate(
                values,
                start=1
            ):

                cell = (
                    findings_sheet.cell(
                        row=row,
                        column=column_index
                    )
                )

                cell.value = value

                cell.alignment = (
                    Alignment(
                        wrap_text=True,
                        vertical="top"
                    )
                )

            row += 1

    #
    # Column widths
    #

    widths = {

        1: 25,
        2: 18,
        3: 35,
        4: 35,
        5: 18,
        6: 35,
        7: 50,
        8: 45,
        9: 15

    }

    for column_index, width in widths.items():

        findings_sheet.column_dimensions[
            get_column_letter(
                column_index
            )
        ].width = width

    findings_sheet.freeze_panes = "A2"

    #
    # Return downloadable bytes
    #

    output = io.BytesIO()

    workbook.save(
        output
    )

    output.seek(0)

    return output.getvalue()