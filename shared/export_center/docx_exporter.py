from io import BytesIO

from docx import Document


def export_docx(
    title,
    content
):
    """
    Export Microsoft Word document while preserving
    Markdown heading levels.
    """

    document = Document()

    #
    # Document Title
    #

    document.add_heading(
        title,
        level=1
    )

    #
    # Process Markdown
    #

    for line in content.splitlines():

        line = line.strip()

        if not line:
            continue

        #
        # Heading 1
        #

        if line.startswith("# "):

            document.add_heading(
                line[2:].strip(),
                level=1
            )

            continue

        #
        # Heading 2
        #

        if line.startswith("## "):

            document.add_heading(
                line[3:].strip(),
                level=2
            )

            continue

        #
        # Heading 3
        #

        if line.startswith("### "):

            document.add_heading(
                line[4:].strip(),
                level=3
            )

            continue

        #
        # Heading 4
        #

        if line.startswith("#### "):

            document.add_heading(
                line[5:].strip(),
                level=4
            )

            continue

        #
        # Heading 5
        #

        if line.startswith("##### "):

            document.add_heading(
                line[6:].strip(),
                level=5
            )

            continue

        #
        # Bullet Lists
        #

        if line.startswith("- "):

            document.add_paragraph(
                line[2:].strip(),
                style="List Bullet"
            )

            continue

        if line.startswith("* "):

            document.add_paragraph(
                line[2:].strip(),
                style="List Bullet"
            )

            continue

        #
        # Numbered Lists
        #

        if line[:2].isdigit() and line[2:4] == ". ":

            document.add_paragraph(
                line[4:].strip(),
                style="List Number"
            )

            continue

        #
        # Bold Text Cleanup
        #

        clean_line = line.replace(
            "**",
            ""
        )

        #
        # Normal Paragraph
        #

        document.add_paragraph(
            clean_line
        )

    output = BytesIO()

    document.save(
        output
    )

    output.seek(
        0
    )

    return output